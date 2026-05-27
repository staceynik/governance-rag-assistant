import os
import uuid
import json
import re
from typing import List, Optional, Dict, Tuple
from dotenv import load_dotenv

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

from .schemas import (
    AskResponse,
    ComplianceVerdict,
)

load_dotenv()

def build_embeddings():

    backend = os.getenv("LLM_BACKEND", "disabled")

    if backend == "openai":
        from langchain_openai import OpenAIEmbeddings

        return OpenAIEmbeddings(
            model=os.getenv(
                "OPENAI_EMBED_MODEL",
                "text-embedding-3-small"
            )
        )

    from langchain_community.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(
        model_name=os.getenv(
            "HF_EMBED_MODEL",
            "sentence-transformers/all-MiniLM-L6-v2"
        )
    )


def build_llm():

    backend = os.getenv("LLM_BACKEND", "disabled")

    if backend == "openai":
        from langchain_openai import ChatOpenAI

        return ChatOpenAI(
            model=os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
            temperature=0
        )

    elif backend == "ollama":
        from langchain_community.chat_models import ChatOllama

        return ChatOllama(
            model="phi3:mini",
            temperature=0
        )

    return None

# -----------------------------
# Simple loaders (txt/pdf/docx)
# -----------------------------
def load_text_from_file(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or "")
        return "\n".join(pages).strip()

    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        import io
        f = io.BytesIO(content)
        doc = DocxDocument(f)
        return "\n".join([p.text for p in doc.paragraphs]).strip()

    # fallback
    return content.decode("utf-8", errors="ignore")

# -----------------------------
# Documents storage + RAG index
# -----------------------------
class RAGStore:
    def __init__(self):
        self.embeddings = build_embeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=int(os.getenv("CHUNK_SIZE", "900")),
            chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "150")),
        )
        self.vectorstore: Optional[FAISS] = None

        # doc_id -> (doc_type, raw_text)
        self.docs: Dict[str, Tuple[str, str]] = {}
        self.storage_path = "storage"
        self.index_path = os.path.join(self.storage_path, "faiss_index")
        self.docs_path = os.path.join(self.storage_path, "docs.json")

        os.makedirs(self.storage_path, exist_ok=True)

        self._load_storage()

    def _save_storage(self):

        if self.vectorstore is not None:
            self.vectorstore.save_local(self.index_path)

        with open(self.docs_path, "w") as f:
            json.dump(self.docs, f)


    def _load_storage(self):

        if os.path.exists(self.docs_path):
            with open(self.docs_path, "r") as f:
                raw_docs = json.load(f)

            self.docs = {
                k: tuple(v)
                for k, v in raw_docs.items()
             }

        if os.path.exists(self.index_path):
            self.vectorstore = FAISS.load_local(
                self.index_path,
                self.embeddings,
                allow_dangerous_deserialization=True
            )

    def _rebuild_index(self) -> int:
        all_docs: List[Document] = []
        for doc_id, (doc_type, text) in self.docs.items():
            all_docs.append(Document(
                page_content=text,
                metadata={
                    "doc_id": doc_id,
                    "doc_type": doc_type,
                },
            ))

        if not all_docs:
            self.vectorstore = None
            return 0

        chunks = self.text_splitter.split_documents(all_docs)
        for idx, chunk in enumerate(chunks):
            chunk.metadata["chunk_id"] = idx
        self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
        return len(chunks)

    def upsert_document(
        self,
        *,
        doc_type: str,
        filename: str,
        content: bytes,
        doc_id: Optional[str] = None
    ) -> Tuple[str, int]:

        text = load_text_from_file(filename, content)

        if not text.strip():
            raise ValueError(
                "File was parsed successfully, but no readable text was found."
            )

        if doc_id is None:
            clean_name = re.sub(r"[^a-zA-Z0-9]+", "_", filename.lower())
            clean_name = clean_name.replace(".pdf", "").replace(".md", "")
            clean_name = clean_name.strip("_")

            doc_type_str = str(doc_type).replace("DocType.", "")

            doc_id = f"{doc_type_str}_{clean_name}"

        self.docs[doc_id] = (doc_type, text)

        chunks_indexed = self._rebuild_index()

        self._save_storage()

        return doc_id, chunks_indexed

    def list_docs(self):
        docs = []

        for doc_id, content in self.docs.items():
            docs.append({
                "doc_id": doc_id,
                "chars": len(content[1]),
            })

        return docs

    def _search(self, query: str, k: int = 8, doc_ids: Optional[List[str]] = None, doc_type: Optional[str] = None) -> List[Document]:
        if self.vectorstore is None:
            return []

        docs = self.vectorstore.similarity_search(query, k=k * 2)

        def ok(d: Document) -> bool:
            md = d.metadata or {}
            if doc_ids is not None and md.get("doc_id") not in set(doc_ids):
                return False
            if doc_type is not None and md.get("doc_type") != doc_type:
                return False
            return True

        filtered = [d for d in docs if ok(d)]
        unique = []
        seen = set()

        for d in filtered:
            text = d.page_content[:200]

            if text not in seen:
                seen.add(text)
                unique.append(d)

        return unique[:k]

    def ask(self, question: str, doc_ids: Optional[List[str]] = None) -> Tuple[str, List[str]]:
        llm = build_llm()

        ctx_docs = self._search(question, k=3, doc_ids=doc_ids)
        context = "\n\n---\n\n".join([f"[{d.metadata.get('doc_type')}:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in ctx_docs])
        context = context[:4000]

        prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                "You are an HR document assistant. "
                "Use only the provided context. "
                "If the answer cannot be derived from the context, explicitly say so. "
                "Do not infer missing legal or HR information. "
                "Keep answers concise and factual."
            ),
            (
                "human",
                "Context:\n{context}\n\n"
                "Question:\n{question}\n\n"
                "Answer:"
            ),
        ])

        msg = prompt.format_messages(context=context, question=question)

        if llm is None:
            answer = (
                "LLM disabled. Retrieved context:\n\n"
                + context[:3000]
            )
        else:
            try:
                out = llm.invoke(msg)
                answer = out.content
            except Exception as e:
                answer = f"LLM generation failed: {str(e)}"

        citations = []
        for d in ctx_docs:
            md = d.metadata or {}
            citations.append(f"{md.get('doc_type')}:{md.get('doc_id')}")

        return answer, sorted(set(citations))

    def check_compliance(self, doc_ids: Optional[List[str]] = None, focus: str = "vacations") -> ComplianceVerdict:
        llm = build_llm()
        if llm is None:
            return ComplianceVerdict(
                compliant=False,
                summary="LLM backend is disabled.",
                violations=[],
                evidence=[],
                recommendations=[
                    "Enable Ollama or OpenAI backend for compliance analysis."
                ]
            )

        # Retrieve policy documents and supporting documents separately
        policy_docs = self._search(f"company policy {focus} rules requirements", k=10, doc_ids=doc_ids, doc_type="policy")
        supporting_docs  = self._search(f"{focus} vacation minimum duration carry over payment schedule", k=14, doc_ids=doc_ids)

        policy_ctx = "\n\n---\n\n".join([f"[policy:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in policy_docs])
        docs_ctx   = "\n\n---\n\n".join([f"[{d.metadata.get('doc_type')}:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in supporting_docs])

        verdict_chain = llm.with_structured_output(ComplianceVerdict)

        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "You are an HR compliance analyst. "
             "Your task is to verify whether the provided documents comply with the company's vacation policy.\n"
             "Rules:\n"
             "- Do not invent facts.\n"
             "- If information is missing, explicitly mention it in the summary and recommendations.\n"
             "- In evidence, provide short quotations (1-3 lines) and reference the source as [policy:ID] or [contract:ID].\n"),
            ("human",
            "Policy excerpts:\n{policy_ctx}\n\n"
            "Document excerpts:\n{docs_ctx}\n\n"
            "Check compliance for the following focus area: {focus}\n"
            "Return a structured compliance verdict."),
        ])

        msg = prompt.format_messages(policy_ctx=policy_ctx, docs_ctx=docs_ctx, focus=focus)
        try:
            return verdict_chain.invoke(msg)
        except Exception as e:
            return ComplianceVerdict(
                compliant=False,
                summary=f"Compliance analysis failed: {str(e)}",
                violations=[],
                evidence=[],
                recommendations=["Retry with fewer documents or a different backend."]
            )
            